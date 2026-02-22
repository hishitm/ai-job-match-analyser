"""
Resume parser: extracts raw text from PDF and DOCX files.
"""

import io
import re
import pdfplumber
from docx import Document


def clean_text(text: str) -> str:
    """Remove excessive whitespace and non-printable characters."""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\x20-\x7E\n]', ' ', text)
    return text.strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pdfplumber."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return clean_text("\n".join(text_parts))


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    # Also grab text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return clean_text("\n".join(paragraphs))


def parse_resume(file_bytes: bytes, filename: str) -> dict:
    """
    Parse a resume file and return extracted text and metadata.
    Returns:
        {
            "text": str,          # Full extracted text
            "filename": str,
            "file_type": str,     # "pdf" or "docx"
            "word_count": int,
            "char_count": int
        }
    """
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
        file_type = "pdf"
    elif filename_lower.endswith((".docx", ".doc")):
        text = extract_text_from_docx(file_bytes)
        file_type = "docx"
    else:
        raise ValueError(f"Unsupported file type: {filename}. Please upload PDF or DOCX.")

    if not text or len(text) < 50:
        raise ValueError("Could not extract meaningful text from the file. Ensure the file is not scanned/image-only.")

    return {
        "text": text,
        "filename": filename,
        "file_type": file_type,
        "word_count": len(text.split()),
        "char_count": len(text)
    }
